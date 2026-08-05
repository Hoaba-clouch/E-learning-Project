from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SYSTEM_ANALYSIS_DATABASE.md"
OUTPUT = ROOT / "Tai_lieu_Tinh_nang_User_Story_LearnX_FINAL.docx"

BLUE = "0B5CAD"
LIGHT_BLUE = "DCEAF7"
PALE_BLUE = "EEF5FB"
GRAY = "666666"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
RED = "FCE4D6"


ACTORS = {
    "G": "khách truy cập",
    "S": "học viên",
    "I": "giảng viên",
    "A": "quản trị viên",
    "P": "hệ thống thanh toán VNPAY",
}

ACTOR_TITLES = {
    "G": "Khách truy cập (Guest)",
    "S": "Học viên (Student)",
    "I": "Giảng viên (Teacher/Instructor)",
    "A": "Quản trị viên (Admin)",
    "P": "Hệ thống ngoài VNPAY",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 8.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths: list[float], font_size: float = 8.0) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)
        if row_index == 0:
            repeat_table_header(row)
            for cell in row.cells:
                set_cell_shading(cell, BLUE)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        elif row_index % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, PALE_BLUE)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instr_text)
    run._r.append(field_separate)
    run._r.append(field_end)


def parse_user_stories() -> list[dict[str, str]]:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    current_actor = ""
    current_group = ""
    stories: list[dict[str, str]] = [
        {
            "code": "UC-G01",
            "actor": "G",
            "group": "Khám phá khóa học",
            "feature": "Xem danh sách khóa học công khai",
            "description": "Xem các khóa APPROVED còn ít nhất một lớp nhận đăng ký.",
        },
        {
            "code": "UC-G02",
            "actor": "G",
            "group": "Khám phá khóa học",
            "feature": "Xem chi tiết khóa học",
            "description": "Xem mô tả, chương trình, lớp, lịch học, học phí và đánh giá.",
        },
        {
            "code": "UC-G03",
            "actor": "G",
            "group": "Tài khoản",
            "feature": "Đăng ký tài khoản học viên",
            "description": "Tạo tài khoản STUDENT với email và số điện thoại không trùng.",
        },
    ]

    section_map = {
        "## 3.3": "S",
        "## 3.4": "I",
        "## 3.5": "A",
        "## 3.6": "P",
    }
    row_pattern = re.compile(r"^\|\s*(UC-([SIAP])\d+)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|$")

    for line in lines:
        for prefix, actor in section_map.items():
            if line.startswith(prefix):
                current_actor = actor
                current_group = ""
                break

        if line.startswith("### Nhóm"):
            current_group = line.split("-", 1)[-1].strip()
        elif line.startswith("## 3.6"):
            current_group = "Thanh toán ngoài hệ thống"

        match = row_pattern.match(line)
        if not match:
            continue

        code, actor, feature, description = match.groups()
        if actor != current_actor:
            continue
        stories.append(
            {
                "code": code.strip(),
                "actor": actor,
                "group": current_group or "Chức năng nghiệp vụ",
                "feature": feature.strip(),
                "description": description.strip(),
            }
        )

    expected = {"G": 3, "S": 43, "I": 40, "A": 15, "P": 4}
    actual = Counter(story["actor"] for story in stories)
    if dict(actual) != expected:
        raise RuntimeError(f"Danh mục user story không đủ: expected={expected}, actual={dict(actual)}")
    return stories


def user_story_text(story: dict[str, str]) -> str:
    actor = ACTORS[story["actor"]]
    feature = story["feature"].lower()

    if feature.startswith("đăng nhập"):
        benefit = "truy cập an toàn những chức năng đúng với vai trò của mình"
    elif feature.startswith("đăng xuất"):
        benefit = "kết thúc phiên và bảo vệ tài khoản khi không còn sử dụng"
    elif "ghi nhớ đăng nhập" in feature:
        benefit = "không phải nhập lại thông tin trong thời hạn phiên được cho phép"
    elif feature.startswith("xem"):
        benefit = "có thông tin cần thiết để học tập hoặc thực hiện bước tiếp theo"
    elif feature.startswith(("tìm", "lọc")):
        benefit = "nhanh chóng tìm đúng dữ liệu mình cần"
    elif feature.startswith(("tạo", "mở", "bắt đầu")):
        benefit = "khởi tạo công việc đúng quy trình nghiệp vụ"
    elif feature.startswith(("sửa", "cập nhật", "chỉnh sửa", "đổi")):
        benefit = "duy trì dữ liệu chính xác và phù hợp với tình trạng hiện tại"
    elif feature.startswith("xóa"):
        benefit = "loại bỏ dữ liệu không còn sử dụng mà không ảnh hưởng dữ liệu ngoài quyền"
    elif feature.startswith(("thêm", "chọn")):
        benefit = "chuẩn bị đúng nội dung trước khi hoàn tất giao dịch"
    elif feature.startswith(("thanh toán", "nhận kết quả")):
        benefit = "hoàn tất giao dịch và được ghi danh chính xác"
    elif feature.startswith(("nộp", "trả lời", "lưu nháp")):
        benefit = "hoàn thành hoạt động đánh giá và không bị mất dữ liệu"
    elif feature.startswith("chấm"):
        benefit = "ghi nhận kết quả học tập minh bạch và có phản hồi"
    elif feature.startswith(("bình luận", "phản hồi", "thả", "đánh giá", "báo cáo")):
        benefit = "tương tác an toàn và cải thiện chất lượng học tập"
    elif feature.startswith(("khóa", "duyệt", "từ chối", "ẩn")):
        benefit = "kiểm soát nội dung hoặc tài khoản theo đúng chính sách hệ thống"
    elif feature.startswith(("quản lý", "phân quyền")):
        benefit = "kiểm soát dữ liệu và quyền hạn thuộc trách nhiệm của mình"
    elif feature.startswith("xuất"):
        benefit = "sử dụng dữ liệu trong báo cáo và đối soát ngoài hệ thống"
    else:
        benefit = "hoàn thành nghiệp vụ của mình một cách chính xác và thuận tiện"
    return f"Là {actor}, tôi muốn {feature} để {benefit}."


SPECIAL_ACCEPTANCE = {
    "UC-G03": "Cho trước email/số điện thoại chưa tồn tại; khi nhập đủ dữ liệu hợp lệ; thì tạo một user ACTIVE có role STUDENT. Dữ liệu trùng hoặc mật khẩu không đạt yêu cầu phải bị từ chối.",
    "UC-S01": "Cho trước tài khoản STUDENT ACTIVE; khi nhập đúng email/số điện thoại và mật khẩu; thì nhận session/token và vào Student Portal. Sai thông tin hoặc sai role phải bị từ chối.",
    "UC-S02": "Cho trước học viên chọn Ghi nhớ đăng nhập; khi đăng nhập thành công rồi reload trang; thì phiên còn hiệu lực trong thời hạn được cấu hình và không quay về trang login.",
    "UC-S09": "Cho trước khóa có batch OPEN/STARTED còn hạn và còn chỗ; khi học viên chọn lớp; thì batch được chọn. Lớp hết hạn, đầy hoặc học viên đã ghi danh khóa đó phải bị chặn.",
    "UC-S10": "Cho trước batch hợp lệ; khi thêm vào giỏ; thì giỏ lưu đúng batch và giá tại thời điểm chọn. Không được tồn tại hai batch của cùng một khóa trong giỏ.",
    "UC-S13": "Cho trước giỏ có ít nhất một batch hợp lệ; khi chọn VNPAY; thì backend tạo txnRef duy nhất, amount đúng và payment URL có chữ ký HMAC SHA-512.",
    "UC-S14": "Cho trước Return URL có chữ ký hợp lệ và response code 00; khi backend xác minh; thì payment chuyển SUCCESS, tạo đúng một enrollment ACTIVE và xóa item đã thanh toán. Callback lặp không tạo dữ liệu trùng.",
    "UC-S19": "Cho trước học viên đã ACTIVE trong khóa; khi hoàn thành bài; thì tiến độ được lưu cho đúng user/lesson và phần trăm khóa được tính lại. Người ngoài lớp không được cập nhật.",
    "UC-S29": "Cho trước attempt IN_PROGRESS và còn thời gian; khi học viên lưu câu trả lời; thì đáp án mới nhất được giữ lại. Attempt đã nộp hoặc hết thời gian không được sửa.",
    "UC-S30": "Cho trước attempt hợp lệ; khi học viên nộp; thì trạng thái chuyển SUBMITTED/COMPLETED, thời gian nộp được lưu và điểm tự động được tính với câu hỗ trợ tự chấm.",
    "UC-S33": "Cho trước bài tập thuộc lớp đã ghi danh và còn quyền nộp; khi gửi file/nội dung/link hợp lệ; thì tạo hoặc cập nhật một submission của học viên.",
    "UC-S41": "Cho trước học viên đủ điều kiện ghi danh/thanh toán/tiến độ; khi gửi rating 1–5 và bình luận tối đa 2.000 ký tự; thì tạo một review. Người chưa đủ điều kiện nhận 403.",
    "UC-I06": "Cho trước giảng viên đã đăng nhập; khi nhập dữ liệu khóa hợp lệ; thì tạo khóa DRAFT gắn đúng teacher_id. Giảng viên khác không được sửa khóa này.",
    "UC-I09": "Cho trước khóa DRAFT thuộc giảng viên và có nội dung cần thiết; khi gửi duyệt; thì trạng thái chuyển PENDING và khóa chờ Admin xử lý.",
    "UC-I19": "Cho trước khóa thuộc sở hữu giảng viên; khi tạo lớp với ngày, sĩ số và học phí hợp lệ; thì batch được tạo. max_students phải dương và thời gian đăng ký hợp lệ.",
    "UC-I22": "Cho trước batch và quy tắc lịch hợp lệ; khi sinh lịch định kỳ; thì tạo các session trong khoảng ngày lớp, bỏ qua buổi trùng và báo số buổi đã tạo/bỏ qua.",
    "UC-I24": "Cho trước session thuộc lớp của giảng viên; khi lưu điểm danh; thì mỗi học viên nhận trạng thái PRESENT/LATE/EXCUSED/ABSENT và dữ liệu ngoài lớp bị từ chối.",
    "UC-I33": "Cho trước submission thuộc assignment/khóa của giảng viên; khi nhập điểm trong giới hạn và phản hồi; thì kết quả được lưu và học viên nhìn thấy.",
    "UC-I38": "Cho trước review của khóa thuộc giảng viên; khi gửi phản hồi; thì teacher_comment được lưu và học viên nhận/xem thông báo liên quan.",
    "UC-A03": "Cho trước Admin có quyền users; khi khóa/mở học viên; thì status đổi đúng ACTIVE/SUSPENDED/INACTIVE. Admin thiếu quyền phải nhận 403.",
    "UC-A09": "Cho trước khóa PENDING và Admin có quyền courses; khi duyệt; thì trạng thái chuyển APPROVED và khóa có thể xuất hiện ở catalog khi có batch hợp lệ.",
    "UC-A10": "Cho trước khóa PENDING; khi từ chối; thì trạng thái chuyển REJECTED và lưu được lý do để giảng viên biết cần sửa gì.",
    "UC-P04": "Cho trước dữ liệu Return từ VNPAY; khi backend tính lại HMAC; thì chỉ payload có secure hash khớp mới được ghi nhận. Payload bị sửa phải bị từ chối và không tạo enrollment.",
}


def acceptance_text(story: dict[str, str]) -> str:
    if story["code"] in SPECIAL_ACCEPTANCE:
        return SPECIAL_ACCEPTANCE[story["code"]]

    actor = story["actor"]
    preconditions = {
        "G": "website và API đang hoạt động",
        "S": "học viên ACTIVE đã đăng nhập và có quyền với dữ liệu liên quan",
        "I": "giảng viên ACTIVE đã đăng nhập và dữ liệu thuộc khóa/lớp mình sở hữu",
        "A": "Admin đã đăng nhập và có permission phù hợp",
        "P": "cấu hình merchant, txnRef và chữ ký đúng quy ước",
    }
    return (
        f"Cho trước {preconditions[actor]}; khi thực hiện “{story['feature']}”; "
        f"thì hệ thống {story['description'][0].lower() + story['description'][1:]} "
        "Yêu cầu sai quyền hoặc dữ liệu không hợp lệ phải bị từ chối và không làm sai lệch dữ liệu."
    )


def implementation_status(code: str) -> tuple[str, str]:
    partial = {
        "UC-I18": "Một phần: bài học có tài nguyên để đọc/tải; màn quản trị tài nguyên riêng chưa tách rõ.",
        "UC-A06": "Một phần: backend kiểm tra permission; chưa có màn hình/API đầy đủ để cấp lại permission Admin.",
    }
    dependent_data = {
        "UC-S06", "UC-S18", "UC-S22", "UC-S23", "UC-S24", "UC-S25", "UC-S26",
        "UC-S27", "UC-S28", "UC-S29", "UC-S30", "UC-S31", "UC-S32", "UC-S33", "UC-S34",
    }
    external = {"UC-S13", "UC-S14", "UC-P01", "UC-P02", "UC-P03", "UC-P04"}
    if code in partial:
        return "MỘT PHẦN", partial[code]
    if code in external:
        return "CÓ ĐIỀU KIỆN", "Logic local đã có; giao dịch VNPAY ngoài đang bị chặn do merchant chưa được phê duyệt."
    if code in dependent_data:
        return "PHỤ THUỘC DỮ LIỆU", "Có UI/API; cần khóa/lớp/nội dung tương ứng trong MySQL để trình diễn đầy đủ."
    return "ĐÃ CÓ", "Đã tìm thấy luồng UI/API hoặc service tương ứng trong code hiện tại."


CRITICAL_SCENARIOS = [
    ("CR-01", "Giữ phiên đăng nhập", "Học viên đăng nhập với remember=true", "Reload hoặc mở lại route Student", "Vẫn ở Student Portal khi token/session còn hạn", "ECP + E2E + Regression"),
    ("CR-02", "Phân quyền theo vai trò", "Student/Teacher/Admin có token hợp lệ", "Truy cập API hoặc route của vai trò khác", "401/403 hoặc chuyển đúng trang login; không lộ dữ liệu", "Negative + RBAC"),
    ("CR-03", "Lọc batch được mua", "Khóa APPROVED có nhiều batch", "Mở catalog/chi tiết khóa", "Chỉ batch còn hạn, đúng trạng thái và còn chỗ được chọn", "Decision table + ECP"),
    ("CR-04", "Một lớp cho mỗi khóa trong giỏ", "Giỏ đã có batch A của khóa X", "Thêm batch B cùng khóa X", "Giỏ thay A bằng B, không giữ đồng thời hai batch", "State transition"),
    ("CR-05", "Thanh toán VNPAY", "Giỏ hợp lệ và cấu hình merchant", "Tạo giao dịch rồi xác minh Return", "Amount/txnRef/chữ ký đúng; SUCCESS tạo enrollment ACTIVE", "Integration + E2E có kiểm soát"),
    ("CR-06", "Callback lặp", "Payment đã SUCCESS", "Gửi lại cùng callback", "Không tạo payment/enrollment trùng", "Idempotency"),
    ("CR-07", "Suất học cuối", "Lớp tối đa 50 và hiện có 49 học viên", "Hai giao dịch cùng xác nhận suất cuối", "Chỉ một enrollment mới; lớp FULL; giao dịch sau bị chặn", "BVA + Concurrency"),
    ("CR-08", "Quyền đánh giá", "Học viên chưa đủ điều kiện khóa học", "Gửi review hợp lệ về hình thức", "Backend trả 403 và không tạo review", "Negative + Authorization"),
    ("CR-09", "Làm bài kiểm tra", "Attempt còn thời gian", "Lưu nháp, reload, rồi nộp", "Đáp án còn nguyên; nộp một lần; không sửa sau submit", "State transition + Boundary time"),
    ("CR-10", "Ownership giảng viên", "Giảng viên A và khóa thuộc B", "A sửa/xóa/chấm dữ liệu khóa của B", "Backend từ chối, không thay đổi DB", "Negative + Ownership"),
    ("CR-11", "Duyệt khóa", "Khóa DRAFT/PENDING", "Teacher gửi duyệt và Admin approve/reject", "Chuyển trạng thái đúng workflow; catalog chỉ nhận APPROVED hợp lệ", "State transition"),
    ("CR-12", "Visual regression", "Có baseline giao diện đăng nhập", "Chạy Playwright sau khi thay đổi CSS", "Sai khác hình ảnh không vượt ngưỡng 1%", "Visual regression"),
]


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Title", 24, BLUE),
        ("Heading 1", 17, BLUE),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, "1F4E79"),
        ("Heading 4", 11, "1F4E79"),
    ]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "LEARNX – TÀI LIỆU TÍNH NĂNG VÀ USER STORY"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_number(section.footer.paragraphs[0])


def add_cover(document: Document) -> None:
    for _ in range(5):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TÀI LIỆU TÍNH NĂNG VÀ USER STORY")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("HỆ THỐNG E-LEARNING LEARNX")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)

    line = document.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("Phục vụ học tập, thuyết trình và kiểm thử môn Đảm bảo chất lượng phần mềm")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    for _ in range(7):
        document.add_paragraph()
    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        "Người thực hiện: Nguyễn Bá Hoà\n"
        "Phiên bản đối chiếu: Mã nguồn LearnX kiểm thử ngày 05/08/2026\n"
        "Ngày cập nhật: 05/08/2026"
    )
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    document.add_page_break()


def add_note(document: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)


def add_intro(document: Document, stories: list[dict[str, str]]) -> None:
    document.add_heading("1. Mục đích và cách sử dụng tài liệu", level=1)
    document.add_paragraph(
        "Tài liệu này giúp thành viên nhóm hiểu toàn bộ phạm vi nghiệp vụ của LearnX trước khi "
        "thuyết trình hoặc thiết kế kiểm thử. Nội dung được đối chiếu từ router Frontend, API "
        "Backend, service và tài liệu phân tích database hiện có; không coi ý tưởng chưa có code "
        "là chức năng đã hoàn thành."
    )
    add_note(
        document,
        "Ba khái niệm phải phân biệt",
        "Tính năng: hệ thống có khả năng gì. User story: người dùng muốn gì và vì sao. "
        "Test case: dữ liệu/bước kiểm tra cụ thể để chứng minh user story thỏa điều kiện chấp nhận. "
        "Một tính năng có trong code không đồng nghĩa đã được kiểm thử đầy đủ.",
    )

    document.add_heading("1.1. Cách đọc một user story", level=2)
    document.add_paragraph("Cấu trúc chuẩn: Là [vai trò], tôi muốn [hành động] để [giá trị nhận được].")
    p = document.add_paragraph()
    p.add_run("Ví dụ: ").bold = True
    p.add_run(
        "Là học viên, tôi muốn thanh toán giỏ hàng qua VNPAY để được ghi danh chính xác vào lớp đã chọn."
    )
    document.add_paragraph(
        "Điều kiện chấp nhận được viết theo Given–When–Then: Cho trước trạng thái ban đầu; khi người "
        "dùng thực hiện hành động; thì hệ thống phải tạo ra kết quả có thể kiểm tra được."
    )

    document.add_heading("1.2. Quy ước trạng thái triển khai", level=2)
    rows = [
        ("Trạng thái", "Ý nghĩa"),
        ("ĐÃ CÓ", "Có luồng UI/API hoặc service tương ứng trong code hiện tại."),
        ("PHỤ THUỘC DỮ LIỆU", "Code đã có nhưng cần dữ liệu MySQL phù hợp để trình diễn."),
        ("CÓ ĐIỀU KIỆN", "Phụ thuộc dịch vụ ngoài hoặc cấu hình môi trường, ví dụ VNPAY."),
        ("MỘT PHẦN", "Có một phần cơ chế nhưng chưa đủ một luồng quản trị hoàn chỉnh."),
    ]
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text, table.rows[0].cells[1].text = rows[0]
    for status, meaning in rows[1:]:
        cells = table.add_row().cells
        cells[0].text = status
        cells[1].text = meaning
    style_table(table, [4.0, 12.0], 9)

    counts = Counter(story["actor"] for story in stories)
    document.add_heading("2. Tổng quan phạm vi hệ thống", level=1)
    document.add_paragraph(
        "LearnX có bốn vai trò người dùng chính và một hệ thống ngoài. Tổng cộng tài liệu thống kê "
        f"{len(stories)} user story nghiệp vụ."
    )
    summary = [
        ("Đối tượng", "Số user story", "Phạm vi chính"),
        (ACTOR_TITLES["G"], str(counts["G"]), "Khám phá khóa học và đăng ký tài khoản"),
        (ACTOR_TITLES["S"], str(counts["S"]), "Mua khóa, học tập, kiểm tra, bài tập, tương tác và tài khoản"),
        (ACTOR_TITLES["I"], str(counts["I"]), "Quản lý khóa/lớp/nội dung, chấm điểm, học viên và phân tích"),
        (ACTOR_TITLES["A"], str(counts["A"]), "Người dùng, duyệt khóa, cấu hình và nội dung chung"),
        (ACTOR_TITLES["P"], str(counts["P"]), "Xử lý và trả kết quả thanh toán Sandbox"),
    ]
    table = document.add_table(rows=1, cols=3)
    for i, value in enumerate(summary[0]):
        table.rows[0].cells[i].text = value
    for row in summary[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, [5.0, 2.5, 8.5], 9)

    document.add_heading("2.1. Luồng nghiệp vụ xuyên suốt", level=2)
    document.add_paragraph(
        "Guest xem catalog → đăng ký/đăng nhập Student → chọn batch → thêm giỏ → thanh toán → "
        "payment SUCCESS → enrollment ACTIVE → học bài/làm bài/nộp bài → đánh giá. Trong khi đó "
        "Teacher tạo và vận hành khóa/lớp; Admin duyệt khóa, quản lý tài khoản và cấu hình hệ thống."
    )


def add_catalog(document: Document, stories: list[dict[str, str]]) -> None:
    document.add_heading("3. Danh mục đầy đủ tính năng và user story", level=1)
    document.add_paragraph(
        "Cột Điều kiện chấp nhận là cơ sở để nhóm chuyển user story thành test case. Cột Trạng thái "
        "phản ánh code hiện tại, không phải tỷ lệ test coverage."
    )

    for actor in ["G", "S", "I", "A", "P"]:
        actor_stories = [story for story in stories if story["actor"] == actor]
        document.add_heading(f"3.{['G', 'S', 'I', 'A', 'P'].index(actor) + 1}. {ACTOR_TITLES[actor]}", level=2)
        groups: list[str] = []
        for story in actor_stories:
            if story["group"] not in groups:
                groups.append(story["group"])

        for group in groups:
            document.add_heading(group, level=3)
            group_stories = [story for story in actor_stories if story["group"] == group]
            table = document.add_table(rows=1, cols=4)
            headers = ["Mã / Tính năng", "User story", "Điều kiện chấp nhận chính", "Trạng thái hiện tại"]
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header

            for story in group_stories:
                status, note = implementation_status(story["code"])
                cells = table.add_row().cells
                set_cell_text(cells[0], f"{story['code']}\n{story['feature']}", bold=True, color="1F4E79", size=7.5)
                set_cell_text(cells[1], user_story_text(story), size=7.5)
                set_cell_text(cells[2], acceptance_text(story), size=7.3)
                set_cell_text(cells[3], f"{status}\n{note}", bold=True, size=7.2)
                if status == "ĐÃ CÓ":
                    set_cell_shading(cells[3], GREEN)
                elif status in {"PHỤ THUỘC DỮ LIỆU", "CÓ ĐIỀU KIỆN"}:
                    set_cell_shading(cells[3], AMBER)
                else:
                    set_cell_shading(cells[3], RED)
            style_table(table, [3.0, 4.2, 6.0, 3.2], 7.5)
            document.add_paragraph()


def add_critical_scenarios(document: Document) -> None:
    document.add_heading("4. Các kịch bản trọng yếu cần hiểu để thuyết trình", level=1)
    document.add_paragraph(
        "Đây là các tình huống có rủi ro cao. Khi cô hỏi “nhóm đã test kỹ chưa”, nên giải thích "
        "đầu vào, hành động, kết quả mong đợi và kỹ thuật kiểm thử như bảng dưới."
    )
    table = document.add_table(rows=1, cols=6)
    headers = ["Mã", "Kịch bản", "Cho trước (Given)", "Khi (When)", "Thì (Then)", "Kỹ thuật"]
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
    for scenario in CRITICAL_SCENARIOS:
        cells = table.add_row().cells
        for i, value in enumerate(scenario):
            cells[i].text = value
    style_table(table, [1.3, 2.5, 3.2, 3.0, 4.0, 2.4], 7.5)

    document.add_heading("4.1. Hai khái niệm dễ bị hỏi nhất", level=2)
    p = document.add_paragraph(style="List Bullet")
    p.add_run("Concurrency: ").bold = True
    p.add_run(
        "hai giao dịch cùng xử lý suất cuối. Backend phải khóa dòng batch trước khi đếm và ghi danh để không vượt max_students."
    )
    p = document.add_paragraph(style="List Bullet")
    p.add_run("Idempotency: ").bold = True
    p.add_run(
        "cùng callback được gửi nhiều lần nhưng kết quả cuối không đổi và không tạo payment/enrollment trùng."
    )


def add_quality_mapping(document: Document) -> None:
    document.add_heading("5. Từ user story đến kế hoạch kiểm thử", level=1)
    rows = [
        ("Nhóm chức năng", "Rủi ro chính", "Kỹ thuật phù hợp", "Ví dụ dữ liệu"),
        ("Đăng nhập/RBAC", "Sai role, lộ dữ liệu, mất phiên", "ECP, negative, security, E2E", "Đúng/sai mật khẩu; Student gọi API Admin"),
        ("Catalog/batch", "Hiện lớp hết hạn hoặc đã đầy", "Decision table, boundary date", "OPEN/STARTED/FULL; trước/trong/sau deadline"),
        ("Giỏ hàng", "Hai lớp cùng khóa, giá sai", "State transition, integration", "Giỏ rỗng; thêm A rồi đổi B"),
        ("Sĩ số", "Vượt 50 khi đồng thời", "BVA, concurrency", "0, 1, 25, 49, 50, 51"),
        ("VNPAY", "Sửa chữ ký, callback lặp", "Negative, idempotency, integration", "Hash đúng/sai; response 00/khác 00"),
        ("Bài kiểm tra", "Hết giờ, mất đáp án, nộp lặp", "State transition, boundary time", "Trước/đúng/sau close_at"),
        ("Ownership Teacher", "Sửa khóa của người khác", "Authorization negative", "Teacher A thao tác course của B"),
        ("Admin workflow", "Duyệt sai trạng thái/quyền", "Decision table, RBAC", "DRAFT/PENDING/APPROVED/REJECTED/HIDDEN"),
        ("Giao diện", "Bố cục lệch sau sửa CSS", "Visual regression", "Baseline và sai khác pixel"),
        ("Validation API", "Payload lạ gây 500", "Property-based fuzz", "Chuỗi rỗng/dài/Unicode; rating -10..15"),
    ]
    table = document.add_table(rows=1, cols=4)
    for i, value in enumerate(rows[0]):
        table.rows[0].cells[i].text = value
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, [3.5, 4.0, 4.2, 4.5], 8)

    document.add_heading("5.1. Kết quả kiểm thử hiện tại", level=2)
    results = [
        ("Tầng", "Kết quả", "Ý nghĩa"),
        ("Frontend Vitest", "4/4 PASS", "Một số component/route bảo vệ đã được kiểm tra."),
        ("Backend Vitest + Supertest", "133/133 PASS", "Service, middleware và API trọng yếu đã có test."),
        ("Playwright", "7/7 PASS", "5 luồng nghiệp vụ, 1 fuzz và 1 visual regression."),
        ("Tổng test tự động", "144/144 PASS", "Các trường hợp đã thiết kế đều đạt; không có nghĩa hệ thống hết mọi lỗi."),
        ("Coverage toàn Backend", "35,20% statements", "Chưa đạt mục tiêu 75%; nhiều code Instructor chưa được bao phủ."),
        ("Module thanh toán", "80,41% statements", "Module rủi ro cao đã được ưu tiên kiểm thử."),
        ("VNPAY ngoài", "BLOCKED code 71", "Merchant Sandbox chưa được phê duyệt; callback local có kiểm soát đã PASS."),
    ]
    table = document.add_table(rows=1, cols=3)
    for i, value in enumerate(results[0]):
        table.rows[0].cells[i].text = value
    for row in results[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table, [4.0, 3.5, 8.5], 8.5)

    add_note(
        document,
        "Cách diễn đạt đúng khi thuyết trình",
        "Không nói “144 test PASS nên hệ thống không còn lỗi”. Hãy nói: “144 trường hợp đã thiết kế đều PASS; "
        "coverage toàn Backend hiện 35,20%, vì vậy nhóm vẫn ghi nhận rủi ro ở những chức năng chưa được bao phủ.”",
        AMBER,
    )


def add_known_gaps(document: Document) -> None:
    document.add_heading("6. Phần chưa hoàn chỉnh hoặc chưa kiểm thử đầy đủ", level=1)
    gaps = [
        ("Coverage Backend", "35,20%, chưa đạt mục tiêu 75%; cần ưu tiên Instructor service/routes."),
        ("VNPAY Sandbox bên ngoài", "Cổng trả code 71 vì merchant chưa được phê duyệt; chưa hoàn tất thẻ/OTP thật."),
        ("VNPAY IPN", "Chưa có endpoint server-to-server riêng để đối soát nếu người dùng không quay lại Return URL."),
        ("Session", "Lưu trong RAM; restart Backend làm mất phiên."),
        ("Admin permission", "Có middleware kiểm tra quyền nhưng chưa có luồng UI/API đầy đủ để cấp lại permission Admin."),
        ("Applitools/testRigor", "Chưa chạy SaaS vì chưa có API key/workspace; dùng Playwright visual và kịch bản NLP thay thế local."),
        ("Load/Performance", "Chưa có kiểm thử tải hoặc ngưỡng đáp ứng chính thức."),
        ("CI/CD quality gate", "Các test chưa được bắt buộc chạy tự động trên pipeline trước khi merge."),
        ("Liên kết phụ", "Một số link Điều khoản, Chính sách, Quên mật khẩu hoặc mạng xã hội còn mang tính giao diện."),
    ]
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Hạng mục"
    table.rows[0].cells[1].text = "Hiện trạng trung thực"
    for name, state in gaps:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = state
    style_table(table, [4.5, 11.5], 9)


def add_presentation_guide(document: Document) -> None:
    document.add_heading("7. Kịch bản trình bày ngắn", level=1)
    document.add_paragraph(
        "“LearnX là hệ thống E-learning có bốn vai trò: khách, học viên, giảng viên và quản trị viên. "
        "Chúng em thống kê 105 user story. Luồng trọng yếu bắt đầu từ xem khóa học, chọn lớp, thêm "
        "giỏ, thanh toán, ghi danh, học bài và đánh giá. Teacher chịu trách nhiệm tạo/vận hành khóa "
        "học; Admin duyệt nội dung và quản lý tài khoản. Về kiểm thử, nhóm dùng Vitest, Supertest và "
        "Playwright theo nhiều tầng, kết hợp ECP, BVA, negative testing, concurrency, idempotency, fuzz "
        "và visual regression. Hiện 144 test tự động PASS. Tuy nhiên coverage Backend mới 35,20%, "
        "merchant VNPAY ngoài chưa được phê duyệt và chưa có IPN, nên kết luận là đủ điều kiện demo có "
        "điều kiện chứ chưa đủ mức production.”"
    )

    document.add_heading("7.1. Câu hỏi phản biện thường gặp", level=2)
    qa = [
        ("User story khác test case thế nào?", "User story mô tả nhu cầu và giá trị; acceptance criteria định nghĩa khi nào đạt; test case là dữ liệu và bước cụ thể để kiểm chứng."),
        ("Tại sao test PASS hết vẫn chưa production?", "Vì PASS chỉ áp dụng cho trường hợp đã thiết kế. Coverage thấp, chưa test tải, session RAM và VNPAY/IPN còn hạn chế."),
        ("Tại sao dùng BVA 49/50/51?", "Vì lỗi sĩ số thường nằm sát giới hạn 50; cần chứng minh suất cuối được nhận và yêu cầu vượt giới hạn bị chặn."),
        ("Callback VNPAY gửi hai lần thì sao?", "Backend phải idempotent: lần sau không tạo payment hoặc enrollment thứ hai."),
        ("Hai người mua suất cuối thì sao?", "Backend khóa dòng batch trong transaction, kiểm tra lại sĩ số rồi chỉ cho một giao dịch ghi danh."),
        ("403 khác 401 thế nào?", "401 là chưa xác thực/token sai; 403 là đã xác thực nhưng không đủ role, ownership hoặc điều kiện nghiệp vụ."),
        ("Coverage 35,20% có thấp không?", "Có đối với mục tiêu 75% toàn dự án; nhóm ưu tiên module thanh toán đạt 80,41% và ghi nhận phần Instructor là nợ kiểm thử."),
        ("VNPAY đã test thật chưa?", "Đã test logic ký/xác minh và callback local có kiểm soát. Cổng Sandbox ngoài bị code 71 nên chưa hoàn tất thẻ/OTP thật."),
    ]
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Câu hỏi"
    table.rows[0].cells[1].text = "Câu trả lời ngắn"
    for question, answer in qa:
        cells = table.add_row().cells
        cells[0].text = question
        cells[1].text = answer
    style_table(table, [5.0, 11.0], 9)


def main() -> None:
    stories = parse_user_stories()
    document = Document()
    configure_document(document)
    add_cover(document)
    add_intro(document, stories)
    add_catalog(document, stories)
    add_critical_scenarios(document)
    add_quality_mapping(document)
    add_known_gaps(document)
    add_presentation_guide(document)

    document.core_properties.title = "Tài liệu tính năng và User Story LearnX"
    document.core_properties.subject = "Đảm bảo chất lượng phần mềm"
    document.core_properties.author = "Nguyễn Bá Hoà"
    document.core_properties.comments = "Đối chiếu từ code và tài liệu hệ thống LearnX ngày 05/08/2026."
    document.save(OUTPUT)
    print(OUTPUT.name)
    print(f"stories={len(stories)}")


if __name__ == "__main__":
    main()
